# Note- crypto in python37/lib/site-packages capitalised to Crypto, also inside - util, math, publickey and random.
import os
import tkinter.messagebox
import pyrebase  # Pyrebase 5
import docx  # Used python-docx
from docx.shared import Inches  # Used python-docx
from requests.exceptions import ConnectionError as Ce_error
from socket import gaierror as ga_error
from urllib3.exceptions import MaxRetryError as Max_retry_error


def format_state(val):
    if val == "Not Set":
        val = ""
    return val


def customer_data():
    ask = tkinter.messagebox.askyesno("Atreport", "Has atreport been filed in?")
    if ask:
        with open('C:\\pcdmisw\\pcddata\\isocust.dat', 'r') as f:
            f.readline()
            customer_name = f.readline()
        with open('C:\\pcdmisw\\pcddata\\isomac.dat', 'r') as f:
            f.readline()
            customer_mac = f.readline()
        return customer_name, customer_mac
    else:
        return "Manchester Precision", "Delta 18-50-13, sn:026"


def generate(last_generated_datasave):
    try:
        config = {
            "apiKey": "AIzaSyABg9UcBLEWFfQqFnlHKcQCELetmi3_aZI",
            "authDomain": "breakdown-app-test.firebaseapp.com",
            "databaseURL": "https://breakdown-app-test.firebaseio.com",
            "projectId": "breakdown-app-test",
            "storageBucket": "breakdown-app-test.appspot.com",
            "messagingSenderId": "619535969929",
            "appId": "1:619535969929:web:25472e17094f41e6b52058",
            "measurementId": "G-XHPX03J51L"}

        dependencies = f"C:\\Users\\{os.getlogin()}\\OneDrive - Hexagon\\Deps_xSave\\"
        # set doc as a blank SCS with header only
        doc = docx.Document(dependencies + "Cal_forms\\KP3-STA-FM-094 - Service Check Sheet  Issue 01  02 09 19.docx")

        firebase = pyrebase.initialize_app(config)
        db = firebase.database()
        # Authorise
        # auth = firebase.auth()
        # email = "lee.maloney@hexagon.com"
        # password = "hexmet"
        # user = auth.create_user_with_email_and_password(email, password)

        # Set Data
        # data = {"First Name": "Lee",
        #         "Second Name": "Maloney"}
        # db.child("/Companion_Version_1P8").child("users").set(data)
        # print("data added")

        # Get data
        state_result = db.child("Companion_Version_1P8").child("State").get()
        # result = db.child("service").get()
        state = state_result.val()
        note_result = db.child("Companion_Version_1P8").child("Note").get()
        note = note_result.val()

        n1 = note.get("N1").strip('""')
        n2 = note.get("N2").strip('""')
        n3 = note.get("N3").strip('""')
        n4 = note.get("N4").strip('""')
        n5 = note.get("N5").strip('""')
        n6 = note.get("N6").strip('""')
        n7 = note.get("N7").strip('""')
        n8 = note.get("N8").strip('""')
        n9 = note.get("N9").strip('""')
        n10 = note.get("N10").strip('""')
        n11 = note.get("N11").strip('""')
        n12 = note.get("N12").strip('""')
        n13 = note.get("N13").strip('""')
        n14 = note.get("N14").strip('""')
        n15 = note.get("N15").strip('""')
        n16 = note.get("N16").strip('""')
        n17 = note.get("N17").strip('""')
        n18 = note.get("N18").strip('""')
        n19 = note.get("N19").strip('""')
        n20 = note.get("N20").strip('""')
        n21 = note.get("N21").strip('""')
        n22 = note.get("N22").strip('""')
        n23 = note.get("N23").strip('""')
        n24 = note.get("N24").strip('""')
        n25 = note.get("N25").strip('""')
        n26 = note.get("N26").strip('""')
        n27 = note.get("N27").strip('""')
        n28 = note.get("N28").strip('""')
        n29 = note.get("N29").strip('""')
        n30 = note.get("N30").strip('""')
        n31 = note.get("N31").strip('""')
        n32 = note.get("N32").strip('""')
        n33 = note.get("N33").strip('""')

        s1 = format_state(state.get("S1").strip('""'))
        s2 = format_state(state.get("S2").strip('""'))
        s3 = format_state(state.get("S3").strip('""'))
        s4 = format_state(state.get("S4").strip('""'))
        s5 = format_state(state.get("S5").strip('""'))
        s6 = format_state(state.get("S6").strip('""'))
        s7 = format_state(state.get("S7").strip('""'))
        s8 = format_state(state.get("S8").strip('""'))
        s9 = format_state(state.get("S9").strip('""'))
        s10 = format_state(state.get("S10").strip('""'))
        s11 = format_state(state.get("S11").strip('""'))
        s12 = format_state(state.get("S12").strip('""'))
        s13 = format_state(state.get("S13").strip('""'))
        s14 = format_state(state.get("S14").strip('""'))
        s15 = format_state(state.get("S15").strip('""'))
        s16 = format_state(state.get("S16").strip('""'))
        s17 = format_state(state.get("S17").strip('""'))
        s18 = format_state(state.get("S18").strip('""'))
        s19 = format_state(state.get("S19").strip('""'))
        s20 = format_state(state.get("S20").strip('""'))
        s21 = format_state(state.get("S21").strip('""'))
        s22 = format_state(state.get("S22").strip('""'))
        s23 = format_state(state.get("S23").strip('""'))
        s24 = format_state(state.get("S24").strip('""'))
        s25 = format_state(state.get("S25").strip('""'))
        s26 = format_state(state.get("S26").strip('""'))
        s27 = format_state(state.get("S27").strip('""'))
        s28 = format_state(state.get("S28").strip('""'))
        s29 = format_state(state.get("S29").strip('""'))
        if s29 == "CC or Older":
            n29 = "It is highly recommended that the customer considers the age of the controller and equipment on " \
                  "this CMM. Support for this controller has been withdrawn and as such would leave the customer at " \
                  "risk if the controller fails and parts are not available. "
        if s29 == "DC240":
            n29 = "It is highly recommended that the customer considers the age of the controller and equipment on " \
                  "this CMM. Support for this controller has been withdrawn and as such would leave the customer at " \
                  "risk if the controller fails and parts are not available. "

        notes = ("Customer reminded that machine accuracy may change with temperature.\n"
                 "Machine data files backed up locally in datasave folder and with Hexagon.\n"
                 "Calibration void if machine is moved.\n"
                 "Regular servicing by Hexagon MI Ltd is recommended.\n")

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Cambria'
        font.size = docx.shared.Pt(14)

        records = (
            ('PRE-SERVICE CONDITION CHECK:', ""),
            ('Machine Environment', s1),
            ('Machine Cleanliness', s2),
            ('Machine Functionality', s3),
            ('Visible Damage', s4),
            ('Computer & peripheral functionality', s5),
            ('Port Lock ID/Type', s6),
            ("", ""),
            ('MECHANICAL WORK:', ""),
            ('Remove covers and blinds', s7),
            ("Check for damage and wear	", s8),
            ("Oil or Water contamination of air system", s9),
            ("Clean Air filter system", s10),
            ("Replace Oil removing filter", s11),
            ("Check Air Bearing Lifts", s12),
            ("Check Roller Bearings", s13),
            ("Check drive system", s14),
            ("Check Counterbalance system	", s15),
            ("", ""),
            ("ELECTRONIC WORK:", ""),
            ("Check all cables and connections", s16),
            ("Clean/Change control cabinet filters", s17),
            ("Check control cabinet cooling system", s18),
            ("Check control voltages", s19),
            ("Check/adjust optical reader outputs", s20),
            ("Check/clean motor brushes	", s21),
            ("", ""),
            ("SAFETY SYSTEMS:", ""),
            ("Check Air pressure switch functions", s22),
            ("Check Emergency Stops	", s23),
            ("Check light guards & fences", s24),
            ("", ""),
            ("", ""),
            ("", ""),
            ("PERIPHERAL EQUIPMENT:", ""),
            ("Condition of probe system", s25),
            ("Condition of Computer", s26),
            ("Condition of Printer", s27),
            ("Condition of Hand box", s28),
            ("", ""),
            ("", ""),
        )

        customer, cmm = customer_data()
        # Create table and set heading
        table = doc.add_table(rows=4, cols=2)
        row = table.rows[0]
        row.cells[0].text = 'CUSTOMER:'
        row.cells[1].text = customer
        row = table.rows[2]
        row.cells[0].text = 'MACHINE DETAILS:'
        row.cells[1].text = cmm

        # Create table
        for check, result in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(check)
            row_cells[1].text = str(result)

        # run = doc.add_paragraph().add_run("MECHANICAL WORK:")
        # font = run.font
        # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        # p = doc.add_paragraph('aaa')

        doc.add_paragraph("COMMENTS:")
        doc.add_paragraph("Machine Serviced IAW OEM Service Check Sheet.")
        note_list = [n1, n2, n3, n4, n5, n6, n7, n8, n9, n10, n11, n12, n13, n14, n15, n16, n17, n18, n19, n20, n21,
                     n22, n22, n23, n24, n25, n26, n26, n27, n28, n30, n31, n32, n33, n29]
        for x in note_list:
            if x != "Note Not Set":
                doc.add_paragraph(x)

        doc.add_paragraph(notes)
        doc.add_paragraph("SIGNATURE:")
        doc.add_picture(dependencies + "Cal_forms\\sig.jpg", width=Inches(3.5))
        doc.add_paragraph("L Maloney")
        print(doc.paragraphs)
        print(doc.paragraphs[3].text)
        # print(doc.paragraphs[1].text)

        # doc.save(f"C:/Users/{os.getlogin()}/KP3-STA-FM-094 - Service Check Sheet  Issue 01  02 09 19.docx")

        doc.save(last_generated_datasave + "Certificates and Documents\\"
                                           "KP3-STA-FM-094 - Service Check Sheet  Issue 01  02 09 19.docx")
        print('SCS Program run.')
    except ga_error as e:
        tkinter.messagebox.showerror("Connection Error",
                                     "Unable to connect to cloud database. Check internet connection. Error is: "
                                     "ga_error")
        print(F"Error is: {e}")
    except Ce_error as e:
        tkinter.messagebox.showerror("Connection Error",
                                     "Unable to connect to cloud database. Check internet connection. Error is: "
                                     "Ce_error")
        print(F"Error is: {e}")
    except Max_retry_error as e:
        tkinter.messagebox.showerror("Connection Error",
                                     "Unable to connect to cloud database. Check internet connection. Error is: "
                                     "Max_retry_error")
        print(F"Error is: {e}")
    except Exception as e:
        tkinter.messagebox.showerror("Connection Error",
                                     "Unable to connect to cloud database. Check internet connection. Error is: {e}")
        print(F"Error is: {e}")
